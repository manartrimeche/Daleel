"""
File extractors — PDF (3-tier), DOCX, TXT, images.
Each extractor returns list[dict] with keys: text, page, ocr_used
"""

import re
import json
import logging
from pathlib import Path
from typing import Dict, List, Tuple

from app.processing.text_utils import clean_text, is_text_garbled
from app.processing.ocr import ocr_image_array, ocr_image_file, _check_tesseract

logger = logging.getLogger(__name__)

_OCR_DPI = 150


# ── TXT ──

def extract_txt(path: Path, original_filename: str | None = None) -> List[Dict]:
    encodings = ["utf-8", "utf-8-sig", "windows-1256", "cp1256",
                 "iso-8859-6", "iso-8859-1", "latin-1"]
    for enc in encodings:
        try:
            text = path.read_text(encoding=enc)
            return [{"text": text, "page": 1, "ocr_used": False}]
        except (UnicodeDecodeError, LookupError):
            continue
    text = path.read_bytes().decode("utf-8", errors="replace")
    return [{"text": text, "page": 1, "ocr_used": False}]


def extract_jsonl(path: Path, original_filename: str | None = None) -> List[Dict]:
    """Extract text from JSONL/JSON files; falls back to raw text extraction."""
    raw_pages = extract_txt(path, original_filename)
    if not raw_pages:
        return []

    raw_text = raw_pages[0]["text"]
    lines = [ln.strip() for ln in raw_text.splitlines() if ln.strip()]
    if not lines:
        return []

    extracted_parts: List[str] = []
    for ln in lines:
        try:
            item = json.loads(ln)
        except json.JSONDecodeError:
            # Some .json files may not be JSONL; keep original line as text.
            extracted_parts.append(ln)
            continue

        if isinstance(item, dict):
            text = item.get("text")
            if isinstance(text, str) and text.strip():
                extracted_parts.append(text.strip())
                continue

            # Common fallback fields in exported datasets.
            for key in ("content", "chunk", "body", "message"):
                val = item.get(key)
                if isinstance(val, str) and val.strip():
                    extracted_parts.append(val.strip())
                    break
        elif isinstance(item, str) and item.strip():
            extracted_parts.append(item.strip())

    if not extracted_parts:
        return raw_pages

    merged = "\n\n".join(extracted_parts)
    return [{"text": merged, "page": 1, "ocr_used": False}]


# ── PDF helpers ──

def _extract_page_pdfminer(path: Path, page_num: int) -> str:
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.layout import LAParams
        laparams = LAParams(line_margin=0.5, word_margin=0.1,
                            char_margin=2.0, detect_vertical=False)
        text = extract_text(str(path), page_numbers=[page_num - 1],
                            laparams=laparams)
        return clean_text(text) if text else ""
    except ImportError:
        return ""
    except Exception as e:
        logger.warning("  pdfminer failed on page %s: %s", page_num, e)
        return ""


def _extract_all_pdfminer(path: Path) -> List[Dict]:
    try:
        from pdfminer.high_level import extract_text
        from pdfminer.pdfpage import PDFPage
        from pdfminer.layout import LAParams
        laparams = LAParams(line_margin=0.5, word_margin=0.1,
                            char_margin=2.0, detect_vertical=False)
        with open(str(path), "rb") as f:
            total = len(list(PDFPage.get_pages(f)))
        pages = []
        for page_num in range(1, total + 1):
            text = extract_text(str(path), page_numbers=[page_num - 1],
                                laparams=laparams)
            text = clean_text(text) if text else ""
            if text and len(text.strip()) > 20:
                pages.append({"text": text, "page": page_num, "ocr_used": False})
        return pages
    except ImportError:
        logger.warning("  pdfminer.six not installed")
        return []
    except Exception as e:
        logger.warning("  pdfminer full extraction failed: %s", e)
        return []


def _ocr_single_page(page, hint_arabic: bool) -> Tuple[str, bool]:
    try:
        import fitz
        import numpy as np
        from PIL import Image
        # Higher DPI (300) for Arabic — small dots/diacritics need more detail
        # Use RGB for Arabic (grayscale loses diacritic contrast)
        dpi = 300 if hint_arabic else (200 if _check_tesseract() else _OCR_DPI)
        if hint_arabic:
            pix = page.get_pixmap(dpi=dpi)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
        else:
            pix = page.get_pixmap(dpi=dpi, colorspace=fitz.csGRAY)
            img = Image.frombytes("L", (pix.width, pix.height), pix.samples)
        ocr_text = clean_text(ocr_image_array(np.array(img), hint_arabic=hint_arabic))
        if ocr_text and len(ocr_text.strip()) > 20:
            return ocr_text, True
    except Exception as e:
        logger.warning("  OCR error: %s", e)
    return "", False


def _ocr_full_pdf(path: Path, hint_arabic: bool, total_pages: int) -> List[Dict]:
    try:
        import fitz
    except ImportError:
        return []
    logger.info(f"  Full OCR mode: {total_pages} pages at {_OCR_DPI} DPI")
    pages = []
    with fitz.open(str(path)) as doc:
        for num, page in enumerate(doc, start=1):
            if num % 10 == 1 or num == total_pages:
                logger.info(f"  OCR progress: page {num}/{total_pages}")
            text, _ = _ocr_single_page(page, hint_arabic)
            if text:
                pages.append({"text": text, "page": num, "ocr_used": True})
    logger.info(f"  OCR complete: {len(pages)} pages with content")
    return pages


# ── PDF ──

def extract_pdf(path: Path, original_filename: str | None = None) -> List[Dict]:
    """3-tier: PyMuPDF → pdfminer.six → Tesseract/EasyOCR."""
    try:
        import fitz
    except ImportError:
        logger.error("PyMuPDF not installed.")
        return []

    # Use original filename (not UUID path) to detect Arabic hint
    name_to_check = original_filename or path.name
    hint_arabic = bool(re.search(r"[\u0600-\u06FF]", name_to_check))
    logger.info(f"  hint_arabic={hint_arabic} (from '{name_to_check}')")

    garbled_doc = False
    total_pages = 0
    with fitz.open(str(path)) as doc:
        total_pages = len(doc)
        sample_texts = []
        sample_raws = []
        for i, page in enumerate(doc):
            if i >= 5:
                break
            raw = page.get_text("text")
            t = clean_text(raw)
            if len(t.strip()) > 50:
                sample_texts.append(t)
                sample_raws.append(raw)
        if sample_texts:
            combined = "\n".join(sample_texts[:3])
            combined_raw = "\n".join(sample_raws[:3])
            if is_text_garbled(combined, expect_arabic=hint_arabic,
                               raw_text=combined_raw):
                garbled_doc = True

    if garbled_doc:
        logger.info("  PyMuPDF produced garbled text — sampling pdfminer.six")
        pm_ok = False
        for test_pg in [3, 6]:
            if test_pg <= total_pages:
                pm_text = _extract_page_pdfminer(path, test_pg)
                if pm_text and len(pm_text.strip()) > 30 and not is_text_garbled(pm_text, expect_arabic=hint_arabic):
                    pm_ok = True
                    break
        if pm_ok:
            logger.info("  pdfminer.six works — extracting all pages")
            pm_pages = _extract_all_pdfminer(path)
            if pm_pages:
                logger.info(f"  pdfminer.six extracted {len(pm_pages)} pages")
                return pm_pages
        logger.info("  Falling back to OCR")
        return _ocr_full_pdf(path, hint_arabic, total_pages)

    pages = []
    with fitz.open(str(path)) as doc:
        for num, page in enumerate(doc, start=1):
            raw = page.get_text("text")
            text = clean_text(raw)

            if text.strip() and is_text_garbled(text, expect_arabic=hint_arabic,
                                                 raw_text=raw):
                pm_text = _extract_page_pdfminer(path, num)
                if pm_text and len(pm_text.strip()) > 20 and not is_text_garbled(pm_text, expect_arabic=hint_arabic):
                    text = pm_text
                    ocr = False
                else:
                    logger.info(f"  Page {num}/{total_pages}: OCR fallback (garbled text)")
                    text, ocr = _ocr_single_page(page, hint_arabic)
            else:
                ocr = False

            if text and len(text.strip()) > 5:
                pages.append({"text": text, "page": num, "ocr_used": ocr})
    return pages


# ── DOCX ──

def extract_docx(path: Path, original_filename: str | None = None) -> List[Dict]:
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed.")
        return []

    doc = Document(str(path))
    lines = []
    for para in doc.paragraphs:
        stripped = para.text.strip()
        if stripped:
            lines.append(stripped)
    for table in doc.tables:
        for row in table.rows:
            row_txt = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_txt:
                lines.append(row_txt)
    return [{"text": "\n".join(lines), "page": 1, "ocr_used": False}]


# ── Image ──

def extract_image(path: Path, original_filename: str | None = None) -> List[Dict]:
    name_to_check = original_filename or path.name
    hint_arabic = bool(re.search(r"[\u0600-\u06FF]", name_to_check))
    text = clean_text(ocr_image_file(path, hint_arabic=hint_arabic))
    return [{"text": text, "page": 1, "ocr_used": True}] if text else []


# ── Registry ──

EXTRACTORS = {
    ".txt":  extract_txt,
    ".json": extract_jsonl,
    ".jsonl": extract_jsonl,
    ".pdf":  extract_pdf,
    ".docx": extract_docx,
    ".doc":  extract_docx,
    ".png":  extract_image,
    ".jpg":  extract_image,
    ".jpeg": extract_image,
    ".tiff": extract_image,
    ".tif":  extract_image,
    ".bmp":  extract_image,
    ".gif":  extract_image,
}
